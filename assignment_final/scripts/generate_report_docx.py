"""
Generate report_final_iteration4.docx from report_final_iteration4.md
Uses python-docx with embedded figures.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = Path(__file__).parent.parent
MD_PATH = BASE / "outputs/adult_census_income/report_final_iteration4.md"
OUT_PATH = BASE / "outputs/adult_census_income/report_final_iteration4.docx"
FIG_DIR = BASE / "outputs/adult_census_income/figures"

# Figure number → file path mapping (matches report Figure references)
FIGURES = {
    1:  FIG_DIR / "target_distribution.png",
    2:  FIG_DIR / "numeric_distributions.png",
    3:  FIG_DIR / "missingness_summary.png",
    4:  FIG_DIR / "leakage_suspicion_numeric_corr.png",
    5:  FIG_DIR / "numeric_correlation_heatmap.png",
    6:  FIG_DIR / "bivariate_boxplots_by_target.png",
    7:  FIG_DIR / "outlier_iqr_rates.png",
    8:  FIG_DIR / "outlier_isolation_scatter.png",
    9:  FIG_DIR / "module2/model_comparison_with_tuning.png",
    10: FIG_DIR / "module2/ablation_summary.png",
    11: FIG_DIR / "module2/learning_curve_selected_model.png",
    12: FIG_DIR / "module2/confusion_matrix.png",
    13: FIG_DIR / "module2/pr_curve.png",
    14: FIG_DIR / "module2/roc_curve.png",
    15: FIG_DIR / "module2/calibration_curve.png",
    16: FIG_DIR / "module3/slice_recall_sex_race.png",
    17: FIG_DIR / "module3/slice_error_age_hours.png",
    18: FIG_DIR / "module3/permutation_importance_top15.png",
    19: FIG_DIR / "module3/threshold_stress_test.png",
    20: FIG_DIR / "missingness_mechanism_classification.png",
    21: FIG_DIR / "agent_visual_correction_example.png",
    22: FIG_DIR / "module2/mlp_training_curve.png",
    23: FIG_DIR / "module2/repeated_cv_stability.png",
}

FIGURE_CAPTIONS = {
    1:  "Figure 1. Class distribution — 75.9% ≤50K, 24.1% >50K (imbalance ratio 3.15:1).",
    2:  "Figure 2. Numeric feature distributions showing extreme right-skew in capital.gain and capital.loss.",
    3:  "Figure 3. Missingness summary — occupation (5.66%), workclass (5.64%), native.country (1.79%).",
    4:  "Figure 4. Leakage detection — no numeric feature carries point-biserial correlation ≥ 0.9.",
    5:  "Figure 5. Numeric correlation heatmap — education.num and education highly correlated.",
    6:  "Figure 6. Bivariate boxplots — marital.status and capital.gain are strongest class-separating features.",
    7:  "Figure 7. IQR outlier rates — hours.per.week reaches 27.7%.",
    8:  "Figure 8. Isolation Forest anomaly scatter — 3.0% of records flagged as anomalous.",
    9:  "Figure 9. Model comparison with tuning — tuned HGB leads on weighted F1, AUC-PR, and ROC-AUC.",
    10: "Figure 10. Ablation summary — no_capital_features collapses F1 by 3.1pp.",
    11: "Figure 11. Learning curves — validation F1 plateaus at 0.869 by 14,586 training samples.",
    12: "Figure 12. Confusion matrix at threshold 0.36 — 252 false negatives, 385 false positives.",
    13: "Figure 13. Precision-recall curve — AUC-PR = 0.828.",
    14: "Figure 14. ROC curve — AUC = 0.926.",
    15: "Figure 15. Calibration curve — ECE = 0.014, near-diagonal.",
    16: "Figure 16. Recall by relationship and age subgroup — Own-child (0.333) and 16–25 (0.364) under-recalled.",
    17: "Figure 17. Error rates by subgroup — Wife (31.8%) and Husband (24.8%) highest overall error.",
    18: "Figure 18. Permutation importance — marital.status (0.058) and capital.gain (0.051) top features.",
    19: "Figure 19. Threshold stress test — ±0.05 shift from 0.36 changes recall by ~4–5pp.",
    20: "Figure 20. Missingness mechanism classification map with explicit MCAR/MAR/MNAR heuristic labels.",
    21: "Figure 21. Agent visual correction evidence — raw missing counts replaced with normalized percentages.",
    22: "Figure 22. MLP training dynamics — stable convergence under early stopping.",
    23: "Figure 23. Repeated-CV stability — selected model performance remains stable across 15 folds.",
}

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Style tweaks
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_inline_formatted(doc, raw_text, size=11):
    """Add a paragraph with **bold** and *italic* inline markdown formatting."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Split on **bold** and *italic* patterns
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', raw_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(size)
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(size)
        else:
            if part:
                run = p.add_run(part)
                run.font.size = Pt(size)
    return p


def add_table_from_rows(doc, rows):
    """Add a formatted table from list-of-lists (first row = header)."""
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text.strip()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if p.runs:
                p.runs[0].font.size = Pt(9)
                if i == 0:
                    p.runs[0].bold = True
                    p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                set_cell_bg(cell, '1a1a2e')
            elif i % 2 == 1:
                set_cell_bg(cell, 'f5f5f5')
    return table


def insert_figure(doc, fig_num):
    """Insert figure image and caption."""
    fig_path = FIGURES.get(fig_num)
    if fig_path and fig_path.exists():
        doc.add_picture(str(fig_path), width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(FIGURE_CAPTIONS.get(fig_num, f"Figure {fig_num}."))
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.runs[0]
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        doc.add_paragraph()


def parse_markdown_table(lines):
    """Parse markdown table lines → list of lists."""
    rows = []
    for line in lines:
        line = line.strip()
        if re.match(r'^\|[-| :]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return rows


# ── figure placement map: after which section heading to insert figures ───────
# Maps a regex pattern in section heading text → list of figure numbers
FIGURE_AFTER_SECTION = {
    r"1\. Introduction":           [1],
    r"EDA-driven design decisions": [2, 3, 20, 4, 5, 6, 7, 8, 21],
    r"Model shortlist, ablation, and tuning protocol": [9],
    r"Fine-tuning and robustness checks": [22, 23],
    r"3\.2 Ablation evidence": [10],
    r"3\.3 Final test performance": [11, 12, 13, 14, 15],
    r"3\.4 Error and failure-mode analysis": [16, 17, 18, 19],
}


# ── main ──────────────────────────────────────────────────────────────────────

def build_docx():
    md_text = MD_PATH.read_text(encoding='utf-8')
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0
    figures_inserted = set()

    def try_insert_figures(heading_text):
        for pattern, fig_nums in FIGURE_AFTER_SECTION.items():
            if re.search(pattern, heading_text):
                for fn in fig_nums:
                    if fn not in figures_inserted:
                        insert_figure(doc, fn)
                        figures_inserted.add(fn)

    while i < len(lines):
        line = lines[i]

        # --- horizontal rule ---
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # --- H1 ---
        m = re.match(r'^# (.+)', line)
        if m:
            add_heading(doc, m.group(1), 1)
            i += 1
            continue

        # --- H2 ---
        m = re.match(r'^## (.+)', line)
        if m:
            heading_text = m.group(1)
            add_heading(doc, heading_text, 2)
            try_insert_figures(heading_text)
            i += 1
            continue

        # --- H3 ---
        m = re.match(r'^### (.+)', line)
        if m:
            heading_text = m.group(1)
            add_heading(doc, heading_text, 3)
            try_insert_figures(heading_text)
            i += 1
            continue

        # --- markdown table ---
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_markdown_table(table_lines)
            if rows:
                add_table_from_rows(doc, rows)
                doc.add_paragraph()
            continue

        # --- blank line ---
        if not line.strip():
            i += 1
            continue

        # --- normal paragraph (may contain **bold** / *italic*) ---
        text = line.strip()
        if text:
            add_inline_formatted(doc, text)
        i += 1

    # Insert any remaining figures not yet placed (at end)
    for fn in sorted(FIGURES.keys()):
        if fn not in figures_inserted:
            insert_figure(doc, fn)
            figures_inserted.add(fn)

    doc.save(str(OUT_PATH))
    print(f"DOCX written to: {OUT_PATH}")
    print(f"File size: {OUT_PATH.stat().st_size / 1024:.1f} KB")
    print(f"Figures embedded: {sorted(figures_inserted)}")


if __name__ == '__main__':
    build_docx()

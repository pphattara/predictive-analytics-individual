"""Generate a two-page agent-log DOCX with concise iteration-level detail."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import List

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


BASE = Path(__file__).parent.parent
LOG_MD_PATH = BASE / "outputs" / "adult_census_income" / "agent_log.md"
OUT_DOCX_PATH = BASE / "outputs" / "adult_census_income" / "agent_log_detailed.docx"


@dataclass
class LogRow:
    row_id: str
    agent_task: str
    output_summary: str
    decision: str
    rationale: str


def parse_sections(md_text: str) -> list[tuple[str, List[LogRow]]]:
    sections: list[tuple[str, List[LogRow]]] = []
    parts = re.split(r"\n(?=## )", md_text.strip())
    for part in parts:
        part = part.strip()
        if not part.startswith("## "):
            continue
        lines = part.splitlines()
        title = lines[0].replace("## ", "").strip()
        table_lines = [ln for ln in lines if ln.strip().startswith("|")]
        if len(table_lines) < 3:
            sections.append((title, []))
            continue

        rows: List[LogRow] = []
        for line in table_lines[2:]:
            cols = [c.strip() for c in line.strip().strip("|").split("|")]
            if len(cols) < 5:
                continue
            rows.append(
                LogRow(
                    row_id=cols[0],
                    agent_task=cols[1],
                    output_summary=cols[2],
                    decision=cols[3],
                    rationale=cols[4],
                )
            )
        sections.append((title, rows))
    return sections


def decision_counts(rows: List[LogRow]) -> tuple[int, int, int]:
    accepted = sum(1 for r in rows if r.decision.startswith("Accepted"))
    modified = sum(1 for r in rows if r.decision.startswith("Modified"))
    rejected = sum(1 for r in rows if r.decision.startswith("Rejected"))
    return accepted, modified, rejected


def first_match(rows: List[LogRow], prefix: str) -> LogRow | None:
    for row in rows:
        if row.decision.startswith(prefix):
            return row
    return None


def iteration_detail(section_title: str, rows: List[LogRow]) -> str:
    accepted, modified, rejected = decision_counts(rows)
    acc = first_match(rows, "Accepted")
    mod = first_match(rows, "Modified")
    rej = first_match(rows, "Rejected")

    parts = [
        f"Accepted {accepted}, modified {modified}, rejected/corrected {rejected}.",
    ]
    if acc:
        parts.append(
            f"Accepted example: '{acc.agent_task}' verified and kept because {acc.rationale.rstrip('.')}.")
    if mod:
        parts.append(
            f"Modified example: '{mod.agent_task}' adjusted for coursework constraints ({mod.rationale.rstrip('.')})."
        )
    if rej:
        parts.append(
            f"Rejected example: '{rej.agent_task}' replaced after risk review ({rej.rationale.rstrip('.')})."
        )
    parts.append(f"Iteration focus: {section_title}.")
    return " ".join(parts)


def build_docx(sections: list[tuple[str, List[LogRow]]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    title = doc.add_heading("Agent Usage Log — Condensed Two-Page Decision Register", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = doc.add_paragraph(
        "Dataset: Adult Census Income | Scope: Iteration 1 Module 1 through Final Submission "
        "| Format: compact iteration-level detail"
    )
    subtitle.runs[0].font.size = Pt(10)
    doc.add_paragraph(
        "This version summarizes decision activity per iteration to keep the log concise while "
        "retaining accepted/modified/rejected evidence and rationale."
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Iteration Section"
    hdr[1].text = "Rows"
    hdr[2].text = "Accepted"
    hdr[3].text = "Modified"
    hdr[4].text = "Rejected"
    hdr[5].text = "Decision and Rationale (Detailed)"

    for section_title, rows in sections:
        accepted, modified, rejected = decision_counts(rows)
        cells = table.add_row().cells
        cells[0].text = section_title
        cells[1].text = str(len(rows))
        cells[2].text = str(accepted)
        cells[3].text = str(modified)
        cells[4].text = str(rejected)
        cells[5].text = iteration_detail(section_title, rows)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    total_rows = sum(len(rows) for _, rows in sections)
    total_accepted = sum(decision_counts(rows)[0] for _, rows in sections)
    total_modified = sum(decision_counts(rows)[1] for _, rows in sections)
    total_rejected = sum(decision_counts(rows)[2] for _, rows in sections)
    doc.add_paragraph(
        f"Totals: {total_rows} logged decisions across {len(sections)} sections | "
        f"Accepted={total_accepted}, Modified={total_modified}, Rejected/Corrected={total_rejected}."
    )

    doc.save(OUT_DOCX_PATH)


def main() -> None:
    md_text = LOG_MD_PATH.read_text(encoding="utf-8")
    sections = parse_sections(md_text)
    build_docx(sections)
    print(f"written: {OUT_DOCX_PATH}")


if __name__ == "__main__":
    main()
